import json
import re
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from PyQt6.QtCore import QBuffer, QObject, QRegularExpression, QSizeF, QThread, QTimer, Qt, pyqtSignal
from PyQt6.QtGui import (
    QBrush,
    QColor,
    QDragEnterEvent,
    QDropEvent,
    QPageSize,
    QPdfWriter,
    QRegularExpressionValidator,
    QTextDocument,
)
from PyQt6.QtWidgets import (
    QHeaderView,
    QHBoxLayout,
    QLabel,
    QLineEdit,
    QMessageBox,
    QPushButton,
    QProgressBar,
    QTableWidget,
    QTableWidgetItem,
    QTextEdit,
    QSplitter,
    QStackedWidget,
    QStyledItemDelegate,
    QVBoxLayout,
    QWidget,
    QWizardPage,
)

from sj_generator.ai.client import LlmClient
from sj_generator.ai.import_questions import _fingerprint_question_obj, import_questions_from_sources
from sj_generator.config import (
    load_deepseek_config,
    load_kimi_config,
    load_qwen_config,
    to_kimi_llm_config,
    to_llm_config,
    to_qwen_llm_config,
)
from sj_generator.io.source_reader import read_source_text
from sj_generator.models import Question
from sj_generator.ui.compare_highlight import compare_highlight_model_styles
from sj_generator.ui.pages.analysis_pages import _commit_draft_questions_to_db
from sj_generator.ui.state import AiSourceFileItem, WizardState, normalize_ai_concurrency
from sj_generator.ui.constants import PAGE_AI_ANALYSIS, PAGE_AI_IMPORT, PAGE_DEDUPE_RESULT, PAGE_IMPORT_SUCCESS
from sj_generator.ui.pdf_preview import DocumentPdfWebView
from sj_generator.ui.styles import rounded_panel_stylesheet


def _sanitize_filename(name: str) -> str:
    s = (name or "").strip()
    s = re.sub(r'[<>:"/\\\\|?*]+', "_", s)
    s = s.replace("\n", " ").replace("\r", " ").replace("\t", " ")
    s = s.strip(" .")
    return s


LEVEL_PATH_RE = re.compile(r"^\d+\.\d+\.\d+$")
LEVEL_PATH_EDIT_RE = QRegularExpression(r"\d*(?:\.\d*(?:\.\d*)?)?")


def _unique_child_dir(parent: Path, base_name: str) -> Path:
    candidate = parent / base_name
    if not candidate.exists():
        return candidate
    i = 2
    while True:
        c = parent / f"{base_name}_{i}"
        if not c.exists():
            return c
        i += 1


def _rename_project(state: WizardState, *, new_name: str) -> bool:
    project_dir = state.project_dir
    repo_path = state.repo_path
    if project_dir is None or repo_path is None:
        return False
    safe = _sanitize_filename(new_name)
    if not safe:
        return False

    parent = project_dir.parent
    target_dir = _unique_child_dir(parent, safe)
    target_name = target_dir.name
    target_repo = target_dir / f"{target_name}.xlsx"

    state.project_dir = target_dir
    state.repo_path = target_repo
    state.project_name_is_placeholder = False
    return True


def _extract_paths_from_drop_event(event: QDropEvent) -> list[Path]:
    md = event.mimeData()
    if not md.hasUrls():
        return []
    out: list[Path] = []
    for u in md.urls():
        if not u.isLocalFile():
            continue
        p = Path(u.toLocalFile())
        if p.exists():
            out.append(p)
    return out


def _merge_paths_text(existing_text: str, paths: list[Path]) -> str:
    existing = [p.strip() for p in (existing_text or "").split(";") if p.strip()]
    seen: set[str] = set()
    merged: list[str] = []
    for s in existing:
        if s not in seen:
            merged.append(s)
            seen.add(s)
    for p in paths:
        s = str(p)
        if s not in seen:
            merged.append(s)
            seen.add(s)
    return "; ".join(merged)


class _LevelPathItemDelegate(QStyledItemDelegate):
    def createEditor(self, parent, option, index):
        editor = QLineEdit(parent)
        editor.setPlaceholderText("例如：3.2.2")
        editor.setValidator(QRegularExpressionValidator(LEVEL_PATH_EDIT_RE, editor))
        return editor


class AiSelectFilesPage(QWizardPage):
    def __init__(self, state: WizardState) -> None:
        super().__init__()
        self._state = state
        self.setTitle("确认导入资料")
        self.setAcceptDrops(True)

        self._files_table = QTableWidget()
        self._files_table.setColumnCount(3)
        self._files_table.setHorizontalHeaderLabels(["名称", "版本", "层级"])
        self._files_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.Stretch)
        self._files_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeMode.ResizeToContents)
        self._files_table.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeMode.Stretch)
        self._files_table.verticalHeader().setVisible(False)
        self._files_table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self._files_table.setSelectionMode(QTableWidget.SelectionMode.SingleSelection)
        self._files_table.setEditTriggers(
            QTableWidget.EditTrigger.DoubleClicked | QTableWidget.EditTrigger.EditKeyPressed
        )
        self._files_table.setItemDelegateForColumn(2, _LevelPathItemDelegate(self._files_table))
        self._files_table.itemSelectionChanged.connect(self._handle_selected_file_changed)
        self._files_table.itemChanged.connect(self._handle_file_table_item_changed)

        left_top_layout = QVBoxLayout()
        left_top_layout.addWidget(self._files_table, 1)
        left_top_panel = QWidget()
        left_top_panel.setLayout(left_top_layout)

        self._import_reminder = QTableWidget()
        self._import_reminder.setColumnCount(3)
        self._import_reminder.setHorizontalHeaderLabels(["名称", "图片", "表格"])
        self._import_reminder.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.Stretch)
        self._import_reminder.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeMode.ResizeToContents)
        self._import_reminder.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeMode.ResizeToContents)
        self._import_reminder.verticalHeader().setVisible(False)
        self._import_reminder.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
        self._import_reminder.setSelectionMode(QTableWidget.SelectionMode.NoSelection)
        self._import_reminder.setFocusPolicy(Qt.FocusPolicy.NoFocus)
        self._import_reminder.setWordWrap(False)
        left_bottom_layout = QVBoxLayout()
        left_bottom_layout.addWidget(self._import_reminder, 1)
        left_bottom_panel = QWidget()
        left_bottom_panel.setLayout(left_bottom_layout)

        left_splitter = QSplitter(Qt.Orientation.Vertical)
        left_splitter.addWidget(left_top_panel)
        left_splitter.addWidget(left_bottom_panel)
        left_splitter.setChildrenCollapsible(False)
        left_splitter.setStretchFactor(0, 3)
        left_splitter.setStretchFactor(1, 2)
        left_splitter.setSizes([300, 180])

        self._preview_placeholder = QTextEdit()
        self._preview_placeholder.setReadOnly(True)
        self._preview_placeholder.setFrameStyle(0)
        self._preview_placeholder.setStyleSheet("background: transparent; border: none;")
        self._preview_placeholder.setPlainText("")
        self._preview_text = QTextEdit()
        self._preview_text.setReadOnly(True)
        self._preview_text.setFrameStyle(0)
        self._preview_text.setStyleSheet("border: none;")
        self._preview_pdf_view = DocumentPdfWebView(self)
        self._preview_stack = QStackedWidget()
        self._preview_stack.addWidget(self._preview_placeholder)
        self._preview_stack.addWidget(self._preview_text)
        self._preview_stack.addWidget(self._preview_pdf_view)
        self._preview_temp_dir = TemporaryDirectory(prefix="sj_doc_preview_")
        self._preview_pdf_path = Path(self._preview_temp_dir.name) / "preview.pdf"

        preview_frame_layout = QVBoxLayout()
        preview_frame_layout.setContentsMargins(0, 0, 0, 0)
        preview_frame_layout.addWidget(self._preview_stack, 1)
        preview_frame = QWidget()
        preview_frame.setStyleSheet(rounded_panel_stylesheet(background="#ffffff"))
        preview_frame.setLayout(preview_frame_layout)

        right_layout = QVBoxLayout()
        right_layout.addWidget(preview_frame, 1)
        right_panel = QWidget()
        right_panel.setLayout(right_layout)

        content_splitter = QSplitter(Qt.Orientation.Horizontal)
        content_splitter.addWidget(left_splitter)
        content_splitter.addWidget(right_panel)
        content_splitter.setChildrenCollapsible(False)
        content_splitter.setStretchFactor(0, 1)
        content_splitter.setStretchFactor(1, 5)
        content_splitter.setSizes([160, 800])

        layout = QVBoxLayout()
        layout.addWidget(content_splitter, 1)
        self.setLayout(layout)

    def initializePage(self) -> None:
        if self._state.ai_source_files_text:
            paths = self._state.ai_source_files or [
                Path(p.strip()) for p in self._state.ai_source_files_text.split(";") if p.strip()
            ]
            self._set_selected_paths(paths)
            self._update_import_reminder(paths)
        else:
            self._set_selected_paths([])
            self._update_import_reminder([])

    def dragEnterEvent(self, event: QDragEnterEvent) -> None:
        if event.mimeData().hasUrls():
            event.acceptProposedAction()
        else:
            event.ignore()

    def dropEvent(self, event: QDropEvent) -> None:
        paths = _extract_paths_from_drop_event(event)
        paths = [p for p in paths if p.suffix.lower() == ".docx"]
        if paths:
            merged = _merge_paths_text(self._serialize_paths_text(), paths)
            merged_paths = [Path(p.strip()) for p in merged.split(";") if p.strip()]
            self._set_selected_paths(merged_paths, selected_path=paths[0])
            self._update_import_reminder(merged_paths)
            event.acceptProposedAction()
        else:
            event.ignore()

    def validatePage(self) -> bool:
        raw = self._serialize_paths_text()
        if not raw:
            QMessageBox.warning(self, "未选择文件", "请选择待处理的资料文件。")
            return False
        paths = [Path(p.strip()) for p in raw.split(";") if p.strip()]
        paths = [p for p in paths if p.exists()]
        if not paths:
            QMessageBox.warning(self, "文件不存在", "请选择存在的资料文件。")
            return False
        invalid_levels = self._find_invalid_level_paths()
        if invalid_levels:
            names = "、".join(invalid_levels)
            QMessageBox.warning(self, "层级格式无效", f"以下文件的层级不是三级数字格式：{names}\n请输入如 3.2.2 的形式。")
            return False
        items = self._collect_table_items()
        level_paths = sorted({item.level_path.strip() for item in items if item.level_path.strip()})
        if not level_paths:
            QMessageBox.warning(self, "未填写层级", "请在“确认导入资料”页填写层级。")
            return False
        if len(level_paths) > 1:
            QMessageBox.warning(
                self,
                "层级不一致",
                "当前导入流程仅支持同一批次写入同一层级，请将本次资料的层级统一后再继续。",
            )
            return False
        self._state.ai_source_files = paths
        self._state.ai_source_files_text = raw
        self._state.ai_source_file_items = items
        self._state.ai_import_level_path = level_paths[0]
        if self._state.project_name_is_placeholder and self._state.project_dir is not None:
            first = paths[0]
            _rename_project(self._state, new_name=first.stem)
        return True

    def nextId(self) -> int:
        return PAGE_AI_IMPORT

    def _serialize_paths_text(self) -> str:
        parts: list[str] = []
        for i in range(self._files_table.rowCount()):
            item = self._files_table.item(i, 0)
            raw = item.data(Qt.ItemDataRole.UserRole)
            if raw:
                parts.append(str(raw))
        return "; ".join(parts)

    def _set_selected_paths(self, paths: list[Path], *, selected_path: Path | None = None) -> None:
        existing_map = {item.path: item for item in self._state.ai_source_file_items}
        self._files_table.blockSignals(True)
        self._files_table.setRowCount(0)
        selected_row = 0
        updated_items: list[AiSourceFileItem] = []
        for index, path in enumerate(paths):
            raw_path = str(path)
            existing = existing_map.get(raw_path, AiSourceFileItem(path=raw_path))
            default_version = existing.version or self._state.preferred_textbook_version
            updated_items.append(
                AiSourceFileItem(path=raw_path, version=default_version, level_path=existing.level_path)
            )
            self._files_table.insertRow(index)
            name_item = QTableWidgetItem(path.name)
            name_item.setToolTip(raw_path)
            name_item.setData(Qt.ItemDataRole.UserRole, raw_path)
            name_item.setFlags(name_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
            version_item = QTableWidgetItem(default_version)
            level_item = QTableWidgetItem(existing.level_path)
            self._files_table.setItem(index, 0, name_item)
            self._files_table.setItem(index, 1, version_item)
            self._files_table.setItem(index, 2, level_item)
            if selected_path is not None and path == selected_path:
                selected_row = index
        self._state.ai_source_file_items = updated_items
        self._files_table.blockSignals(False)
        if self._files_table.rowCount() == 0:
            self._update_preview([])
            return
        self._files_table.selectRow(selected_row)

    def _handle_selected_file_changed(self) -> None:
        path = self._current_selected_path()
        if path is None:
            self._update_preview([])
            return
        self._update_preview([path])

    def _handle_file_table_item_changed(self, item: QTableWidgetItem) -> None:
        row = item.row()
        path_item = self._files_table.item(row, 0)
        if path_item is None:
            return
        raw_path = str(path_item.data(Qt.ItemDataRole.UserRole) or "").strip()
        if not raw_path:
            return
        version_item = self._files_table.item(row, 1)
        level_item = self._files_table.item(row, 2)
        version = (version_item.text() if version_item else "").strip() or self._state.preferred_textbook_version
        if version_item is not None and version_item.text().strip() != version:
            self._files_table.blockSignals(True)
            version_item.setText(version)
            self._files_table.blockSignals(False)
        level_path = (level_item.text() if level_item else "").strip()
        if item.column() == 2 and level_path and not self._is_valid_level_path(level_path):
            self._files_table.blockSignals(True)
            item.setText("")
            self._files_table.blockSignals(False)
            QMessageBox.warning(self, "层级格式无效", "层级只允许输入三级点连接的数字形式，例如 3.2.2。")
            level_path = ""
        self._upsert_source_file_item(raw_path, version=version, level_path=level_path)

    def _current_selected_path(self) -> Path | None:
        row = self._files_table.currentRow()
        if row < 0:
            return None
        item = self._files_table.item(row, 0)
        if item is None:
            return None
        raw = item.data(Qt.ItemDataRole.UserRole)
        if not raw:
            return None
        return Path(str(raw))

    def _collect_table_items(self) -> list[AiSourceFileItem]:
        items: list[AiSourceFileItem] = []
        for row in range(self._files_table.rowCount()):
            name_item = self._files_table.item(row, 0)
            if name_item is None:
                continue
            raw_path = str(name_item.data(Qt.ItemDataRole.UserRole) or "").strip()
            if not raw_path:
                continue
            version = (
                (self._files_table.item(row, 1).text() if self._files_table.item(row, 1) else "").strip()
                or self._state.preferred_textbook_version
            )
            level_path = (self._files_table.item(row, 2).text() if self._files_table.item(row, 2) else "").strip()
            items.append(AiSourceFileItem(path=raw_path, version=version, level_path=level_path))
        return items

    def _upsert_source_file_item(self, raw_path: str, *, version: str, level_path: str) -> None:
        updated = False
        for item in self._state.ai_source_file_items:
            if item.path == raw_path:
                item.version = version
                item.level_path = level_path
                updated = True
                break
        if not updated:
            self._state.ai_source_file_items.append(
                AiSourceFileItem(path=raw_path, version=version, level_path=level_path)
            )

    def _find_invalid_level_paths(self) -> list[str]:
        invalid_names: list[str] = []
        for row in range(self._files_table.rowCount()):
            name_item = self._files_table.item(row, 0)
            level_item = self._files_table.item(row, 2)
            if name_item is None or level_item is None:
                continue
            level_path = level_item.text().strip()
            if level_path and (not self._is_valid_level_path(level_path)):
                invalid_names.append(name_item.text().strip() or f"第{row + 1}行")
        return invalid_names

    def _is_valid_level_path(self, value: str) -> bool:
        return bool(LEVEL_PATH_RE.fullmatch((value or "").strip()))

    def _update_import_reminder(self, paths: list[Path]) -> None:
        if not paths:
            self._import_reminder.clearContents()
            self._import_reminder.setRowCount(0)
            return

        self._import_reminder.setRowCount(len(paths))
        for row, path in enumerate(paths):
            name_item = QTableWidgetItem(path.name)
            name_item.setToolTip(str(path))
            self._import_reminder.setItem(row, 0, name_item)
            has_image, has_table, error_text = self._inspect_docx_content(path)
            if error_text:
                fail_text = f"失败：{error_text}"
                self._import_reminder.setItem(row, 1, self._make_import_check_item(fail_text, "error"))
                self._import_reminder.setItem(row, 2, self._make_import_check_item(fail_text, "error"))
                continue
            self._import_reminder.setItem(
                row,
                1,
                self._make_import_check_item("发现" if has_image else "未发现", "found" if has_image else "clear"),
            )
            self._import_reminder.setItem(
                row,
                2,
                self._make_import_check_item("发现" if has_table else "未发现", "found" if has_table else "clear"),
            )
        self._import_reminder.resizeRowsToContents()

    def _make_import_check_item(self, text: str, state: str) -> QTableWidgetItem:
        item = QTableWidgetItem(text)
        if state == "found":
            item.setForeground(QBrush(QColor(156, 0, 6)))
            item.setBackground(QBrush(QColor(255, 199, 206)))
        elif state == "clear":
            item.setForeground(QBrush(QColor(0, 97, 0)))
            item.setBackground(QBrush(QColor(198, 239, 206)))
        elif state == "error":
            item.setForeground(QBrush(QColor(156, 0, 6)))
            item.setBackground(QBrush(QColor(255, 235, 156)))
        return item

    def _inspect_docx_content(self, path: Path) -> tuple[bool, bool, str]:
        try:
            doc = Document(str(path))
            has_table = len(doc.tables) > 0
            has_image = len(doc.inline_shapes) > 0
            return has_image, has_table, ""
        except Exception as e:
            return False, False, str(e)

    def _update_preview(self, paths: list[Path]) -> None:
        if not paths:
            self._preview_placeholder.setPlainText("")
            self._preview_stack.setCurrentWidget(self._preview_placeholder)
            return

        path = paths[0]
        ext = path.suffix.lower()
        if ext == ".docx":
            try:
                pdf_data = self._build_docx_preview_pdf(path)
                self._preview_pdf_path.write_bytes(pdf_data)
                self._preview_pdf_view.open_pdf(self._preview_pdf_path)
                self._preview_stack.setCurrentWidget(self._preview_pdf_view)
                return
            except Exception as e:
                self._preview_placeholder.setPlainText(f"文档预览区域\n\nWord 预览失败：{e}")
                self._preview_stack.setCurrentWidget(self._preview_placeholder)
                return

        self._preview_placeholder.setPlainText("文档预览区域\n\n当前仅支持 Word 文档预览。")
        self._preview_stack.setCurrentWidget(self._preview_placeholder)

    def _build_docx_preview_pdf(self, path: Path) -> bytes:
        text = read_source_text(path)
        buffer = QBuffer()
        buffer.open(QBuffer.OpenModeFlag.WriteOnly)
        writer = QPdfWriter(buffer)
        writer.setPageSize(QPageSize(QPageSize.PageSizeId.A4))
        writer.setResolution(96)

        doc = QTextDocument()
        doc.setPlainText(text or "(文档内容为空)")
        page_size = writer.pageLayout().paintRectPixels(writer.resolution()).size()
        doc.setPageSize(QSizeF(page_size))
        doc.print(writer)
        buffer.close()
        return bytes(buffer.data())


class AiImportPage(QWizardPage):
    def __init__(self, state: WizardState) -> None:
        super().__init__()
        self._state = state
        self.setTitle("AI 解析详情")

        self._files_edit = QLineEdit()
        self._files_edit.setReadOnly(True)
        self._files_edit.setPlaceholderText("待处理资料文件")

        self._progress = QProgressBar()
        self._progress.setRange(0, 0)
        self._progress.setFormat("正在统计题数…")

        self._stop_btn = QPushButton("停止解析")
        self._stop_btn.clicked.connect(self._stop)
        self._stop_btn.setEnabled(False)

        self._retry_btn = QPushButton("重试")
        self._retry_btn.clicked.connect(self._retry)
        self._retry_btn.setEnabled(False)

        self._status_label = QLabel("")
        self._status_label.setWordWrap(True)

        self._detail_table = QTableWidget()
        self._detail_table.setColumnCount(7)
        self._detail_table.setHorizontalHeaderLabels(
            ["Deepseek耗时", "Kimi耗时", "千问耗时", "Deepseek返回", "Kimi返回", "千问返回", "对比结论"]
        )
        self._detail_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        self._detail_table.verticalHeader().setSectionResizeMode(QHeaderView.ResizeMode.ResizeToContents)
        self._detail_table.setWordWrap(True)
        self._detail_table.setTextElideMode(Qt.TextElideMode.ElideNone)
        self._detail_table.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
        self._detail_table.setRowCount(0)
        self._detail_table.horizontalHeader().sectionResized.connect(
            lambda *_: self._schedule_detail_row_resize()
        )
        self._detail_row_map: dict[int, int] = {}

        row = QHBoxLayout()
        row.addWidget(self._files_edit, 1)
        row.addWidget(self._stop_btn)
        row.addWidget(self._retry_btn)

        layout = QVBoxLayout()
        layout.addLayout(row)
        layout.addWidget(self._status_label)
        layout.addWidget(self._progress)
        layout.addWidget(self._detail_table)
        self.setLayout(layout)
        
        self._items: list[Question] = []
        self._committed = False
        self._last_files_text: str = ""
        self._thread: QThread | None = None
        self._worker: _AiImportWorker | None = None
        self._running = False
        self._stopped = False
        self._finished = False
        self._failed = False
        self._progress_cur = 0
        self._progress_total = 0
        self._accepted_count = 0
        self._skipped_count = 0
        self._phase_text = "准备开始解析…"
        self._detail_text = ""
        self._parallel_text = ""
        self._consistency_text = ""
        self._detail_row_map = {}
        self._compare_secs: dict[int, dict[str, dict[int, int]]] = {}
        self._detail_row_resize_pending = False
        self._deepseek_ready = False
        self._qwen_ready = False
        self._kimi_ready = False
        self._cur_source_name = "-"
        self._cur_question_no = "-"
        self._cur_round_no = "-"

    def initializePage(self) -> None:
        self._files_edit.setText(self._state.ai_source_files_text)
        if self._state.ai_source_files:
            self._cur_source_name = self._state.ai_source_files[0].name
        if self._state.ai_source_files_text and self._state.ai_source_files_text != self._last_files_text:
            self._last_files_text = self._state.ai_source_files_text
            self._detail_table.setRowCount(0)
            self._detail_row_map = {}
            self._compare_secs = {}
            self._items = []
            self._progress.setRange(0, 0)
            self._progress.setValue(0)
            self._reset_progress_meta()
            self._committed = False
            self._running = True
            self._stopped = False
            self._finished = False
            self._failed = False
            self.completeChanged.emit()
            QTimer.singleShot(0, self._start_import)

    def nextId(self) -> int:
        if self._state.dedupe_enabled:
            return PAGE_DEDUPE_RESULT
        return PAGE_AI_ANALYSIS if self._state.analysis_enabled else PAGE_IMPORT_SUCCESS

    def isComplete(self) -> bool:
        if self._running or self._failed:
            return False
        if not (self._finished or self._stopped):
            return False
        return len(self._items) > 0

    def validatePage(self) -> bool:
        if self._running:
            return False
        if not self._items:
            QMessageBox.warning(self, "暂无结果", "当前没有可进入下一步的题目。")
            return False
        self._state.ai_import_questions = list(self._items)
        self._state.draft_questions = list(self._items)
        self._state.dedupe_hits = None
        self._state.reset_db_import()
        if not self._state.dedupe_enabled and not self._state.analysis_enabled:
            return _commit_draft_questions_to_db(self, self._state)
        self.completeChanged.emit()
        return True

    def _start_import(self) -> None:
        if self._thread is not None and self._thread.isRunning():
            return

        paths = self._state.ai_source_files or []
        if not paths:
            QMessageBox.warning(self, "未选择文件", "请先选择待处理的资料文件。")
            return

        cfg = load_deepseek_config()
        kimi_cfg = load_kimi_config()
        qwen_cfg = load_qwen_config()
        self._deepseek_ready = cfg.is_ready()
        self._kimi_ready = kimi_cfg.is_ready()
        self._qwen_ready = qwen_cfg.is_ready()
        self._render_status()
        if not cfg.is_ready():
            QMessageBox.warning(self, "未配置", "DeepSeek 未配置：请先在配置文件中填写 API Key。")
            return

        self._status_label.setText("正在解析，请稍候…")
        self._reset_progress_meta()
        self._phase_text = "正在解析，请稍候…"
        self._render_status()
        self._stop_btn.setEnabled(True)
        self._retry_btn.setEnabled(False)
        self._running = True
        self._stopped = False
        self._finished = False
        self._failed = False
        self.completeChanged.emit()

        thread = QThread(self)
        worker = _AiImportWorker(
            cfg=cfg,
            paths=paths,
            strategy="per_question",
            max_question_workers=normalize_ai_concurrency(self._state.ai_concurrency),
        )
        worker.moveToThread(thread)
        thread.started.connect(worker.run)
        worker.progress.connect(self._on_progress)
        worker.progress_count.connect(self._on_progress_count)
        worker.question.connect(self._on_question)
        worker.compare.connect(self._on_compare)
        worker.done.connect(self._on_done)
        worker.error.connect(self._on_error)
        worker.done.connect(thread.quit)
        worker.error.connect(thread.quit)
        thread.finished.connect(worker.deleteLater)
        thread.finished.connect(thread.deleteLater)

        self._thread = thread
        self._worker = worker
        thread.start()

    def _on_progress(self, msg: str) -> None:
        self._phase_text = msg
        self._update_detail_from_progress(msg)
        self._render_status()

    def _on_progress_count(self, cur: int, total: int) -> None:
        self._progress_cur = max(0, cur)
        self._progress_total = max(0, total)
        if total <= 0:
            self._progress.setRange(0, 0)
            self._progress.setFormat("正在统计题数…")
            self._render_status()
            return
        if self._progress.maximum() != total:
            self._progress.setRange(0, total)
        self._progress.setValue(max(0, min(cur, total)))
        self._progress.setFormat("选择题 %v/%m题")
        self._render_status()

    def _on_question(self, q: Question) -> None:
        self._items.append(q)
        self._accepted_count = len(self._items)
        self._render_status()

    def _on_compare(self, payload: object) -> None:
        if not isinstance(payload, dict):
            return
        idx = int(payload.get("index") or 0)
        if idx <= 0:
            return
        row = self._detail_row_map.get(idx)
        if row is None:
            row = self._detail_table.rowCount()
            self._detail_table.setRowCount(row + 1)
            self._detail_row_map[idx] = row
        round_no = payload.get("round")
        round_no_int = int(round_no or 0)
        self._record_round_sec(
            idx=idx,
            round_no=round_no_int,
            model_key="deepseek",
            sec_value=payload.get("deepseek_sec"),
            ms_value=payload.get("deepseek_ms"),
        )
        self._record_round_sec(
            idx=idx,
            round_no=round_no_int,
            model_key="kimi",
            sec_value=payload.get("kimi_sec"),
            ms_value=payload.get("kimi_ms"),
        )
        self._record_round_sec(
            idx=idx,
            round_no=round_no_int,
            model_key="qwen",
            sec_value=payload.get("qwen_sec"),
            ms_value=payload.get("qwen_ms"),
        )
        self._detail_table.setItem(
            row,
            0,
            QTableWidgetItem(self._format_round_secs(idx=idx, model_key="deepseek")),
        )
        self._detail_table.setItem(
            row,
            1,
            QTableWidgetItem(self._format_round_secs(idx=idx, model_key="kimi")),
        )
        self._detail_table.setItem(
            row,
            2,
            QTableWidgetItem(self._format_round_secs(idx=idx, model_key="qwen")),
        )
        self._detail_table.setItem(row, 3, QTableWidgetItem(self._format_payload_cell(payload.get("deepseek"))))
        self._detail_table.setItem(row, 4, QTableWidgetItem(self._format_payload_cell(payload.get("kimi"))))
        self._detail_table.setItem(row, 5, QTableWidgetItem(self._format_payload_cell(payload.get("qwen"))))
        self._apply_partial_pass_highlight(row=row, payload=payload)
        verdict = self._build_compare_verdict(payload)
        self._detail_table.setItem(row, 6, QTableWidgetItem(verdict))
        self._schedule_detail_row_resize()
        round_no = payload.get("round") or "?"
        self._cur_question_no = str(idx)
        self._cur_round_no = str(round_no)
        self._phase_text = f"解析中：已回传第 {idx} 题第 {round_no}/3 轮比对结果"
        self._render_status()

    def _schedule_detail_row_resize(self) -> None:
        if self._detail_row_resize_pending:
            return
        self._detail_row_resize_pending = True
        QTimer.singleShot(0, self._resize_detail_rows_to_contents)

    def _resize_detail_rows_to_contents(self) -> None:
        self._detail_row_resize_pending = False
        self._detail_table.resizeRowsToContents()

    def _sec_int(self, sec_value: object, ms_value: object) -> int | None:
        if isinstance(sec_value, (int, float)):
            return int(round(float(sec_value)))
        if isinstance(ms_value, (int, float)):
            return int(round(float(ms_value) / 1000.0))
        return None

    def _record_round_sec(
        self,
        *,
        idx: int,
        round_no: int,
        model_key: str,
        sec_value: object,
        ms_value: object,
    ) -> None:
        if idx <= 0 or round_no <= 0:
            return
        sec_int = self._sec_int(sec_value, ms_value)
        if sec_int is None:
            return
        if idx not in self._compare_secs:
            self._compare_secs[idx] = {}
        if model_key not in self._compare_secs[idx]:
            self._compare_secs[idx][model_key] = {}
        self._compare_secs[idx][model_key][round_no] = sec_int

    def _format_round_secs(self, *, idx: int, model_key: str) -> str:
        model_rounds = self._compare_secs.get(idx, {}).get(model_key, {})
        if not model_rounds:
            return ""
        parts = [str(model_rounds[r]) for r in sorted(model_rounds.keys())]
        return "+".join(parts)

    def _clear_compare_cell_bg(self, row: int) -> None:
        for col in [3, 4, 5]:
            item = self._detail_table.item(row, col)
            if item is not None:
                item.setBackground(QBrush())

    def _highlight_sig_text(self, obj: object) -> str:
        if isinstance(obj, dict):
            return _fingerprint_question_obj(obj)
        if obj is None:
            return ""
        if isinstance(obj, str):
            return obj.strip()
        if isinstance(obj, list) and not obj:
            return ""
        return str(obj).strip()

    def _format_payload_cell(self, value: object) -> str:
        if value is None:
            return ""
        try:
            return json.dumps(value, ensure_ascii=False)
        except Exception:
            return str(value)

    def _apply_partial_pass_highlight(self, *, row: int, payload: dict[str, object]) -> None:
        self._clear_compare_cell_bg(row)
        model_cols = [("deepseek", 3), ("kimi", 4), ("qwen", 5)]
        highlight_styles = compare_highlight_model_styles(
            model_sigs={key: self._highlight_sig_text(payload.get(key)) for key, _ in model_cols},
            round_no=int(payload.get("round") or 0),
            round_matched_count=int(payload.get("round_matched_count") or 0),
        )
        if not highlight_styles:
            return
        fail_brush = QBrush(QColor(255, 199, 206))
        empty_brush = QBrush(QColor(255, 235, 156))
        for key, col in model_cols:
            style = highlight_styles.get(key)
            if not style:
                continue
            item = self._detail_table.item(row, col)
            if item is None:
                continue
            if style == "yellow":
                item.setBackground(empty_brush)
            elif style == "red":
                item.setBackground(fail_brush)

    def _build_compare_verdict(self, payload: dict[str, object]) -> str:
        round_no = int(payload.get("round") or 0)
        round_labels = {1: "一轮", 2: "二轮", 3: "三轮"}
        label = round_labels.get(round_no, f"第{round_no}轮")
        round_matched = int(payload.get("round_matched_count") or 0)
        round_valid = int(payload.get("round_valid_count") or 0)
        matched = int(payload.get("matched_count") or 0)
        valid = int(payload.get("valid_count") or 0)
        required = int(payload.get("required_count") or 0)
        accepted = bool(payload.get("accepted"))
        status_text = "通过" if accepted else "未通过"
        if round_no <= 1:
            return f"{label}{matched}/{valid} {status_text}（阈值≥{required}，本轮{round_matched}/{round_valid}）"
        return f"{label}累计{matched}/{valid} {status_text}（阈值≥{required}，本轮{round_matched}/{round_valid}）"

    def _on_done(self, total: int) -> None:
        if self._stopped and not self._failed:
            self._phase_text = f"已停止：当前 {total} 题（可继续进入下一步）"
        else:
            self._phase_text = f"解析完成：{total} 题（可继续进入下一步）"
        if self._progress.maximum() > 0:
            self._progress.setValue(self._progress.maximum())
            self._progress_cur = self._progress.maximum()
            self._progress_total = self._progress.maximum()
        self._stop_btn.setEnabled(False)
        self._retry_btn.setEnabled(False)
        self._running = False
        self._finished = True
        self._thread = None
        self._worker = None
        self._render_status()
        self.completeChanged.emit()

    def _on_error(self, msg: str) -> None:
        QMessageBox.critical(self, "解析失败", msg)
        self._phase_text = "解析失败。"
        self._detail_text = msg
        self._stop_btn.setEnabled(False)
        self._retry_btn.setEnabled(True)
        self._running = False
        self._failed = True
        self._thread = None
        self._worker = None
        self._render_status()
        self.completeChanged.emit()

    def _stop(self) -> None:
        if self._worker is not None:
            self._worker.request_stop()
            self._stop_btn.setEnabled(False)
            self._phase_text = "正在停止…"
            self._stopped = True
            self._render_status()
            self.completeChanged.emit()

    def prepare_to_close(self) -> bool:
        thread = self._thread
        if thread is None or (not thread.isRunning()):
            return True
        if self._worker is not None:
            self._worker.request_stop()
        self._stop_btn.setEnabled(False)
        self._phase_text = "正在停止…"
        self._stopped = True
        self._render_status()
        QMessageBox.information(self, "正在停止", "解析线程仍在收尾，请稍候片刻后再关闭窗口。")
        return False

    def _retry(self) -> None:
        if self._running:
            return
        self._detail_table.setRowCount(0)
        self._detail_row_map = {}
        self._compare_secs = {}
        self._items = []
        self._progress.setRange(0, 0)
        self._progress.setValue(0)
        self._progress.setFormat("正在统计题数…")
        self._reset_progress_meta()
        self._phase_text = "准备重试解析…"
        self._committed = False
        self._stopped = False
        self._finished = False
        self._failed = False
        self._retry_btn.setEnabled(False)
        self._render_status()
        self.completeChanged.emit()
        QTimer.singleShot(0, self._start_import)

    def _reset_progress_meta(self) -> None:
        self._progress_cur = 0
        self._progress_total = 0
        self._accepted_count = 0
        self._skipped_count = 0
        self._detail_text = ""
        self._parallel_text = ""
        self._consistency_text = ""
        self._cur_question_no = "-"
        self._cur_round_no = "-"
        self._render_status()

    def _update_detail_from_progress(self, msg: str) -> None:
        if "已跳过" in msg:
            self._skipped_count += 1
        msrc = re.match(r"^([^：]+)：", msg)
        if msrc:
            self._cur_source_name = msrc.group(1)
        m0 = re.search(r"第\s*(\d+)\s*/\s*(\d+)\s*题（第\s*(\d+)\s*/\s*3\s*轮，三模型并行请求中）", msg)
        if m0:
            i, n, r = m0.groups()
            self._cur_question_no = i
            self._cur_round_no = r
            self._parallel_text = f"并行状态：第 {r}/3 轮三模型并行请求中（题目 {i}/{n}）"
            return
        m1 = re.search(r"第\s*(\d+)\s*/\s*(\d+)\s*题（第\s*(\d+)\s*/\s*3\s*轮，三模型结果(一致|不一致)）", msg)
        if m1:
            i, n, r, st = m1.groups()
            self._consistency_text = f"一致性结论：第 {r}/3 轮{st}（题目 {i}/{n}）"
            if st == "一致":
                self._parallel_text = ""
            return
        m = re.search(r"第\s*(\d+)\s*/\s*(\d+)\s*题（第\s*(\d+)\s*/\s*3\s*轮，([^）]+)）", msg)
        if m:
            i, n, r, model = m.groups()
            self._cur_question_no = i
            self._cur_round_no = r
            self._detail_text = f"当前题：{i}/{n}；轮次：{r}/3；模型：{model}"
            return
        m2 = re.search(r"^(.+?)：统计题数", msg)
        if m2:
            self._cur_source_name = m2.group(1)

    def _render_status(self) -> None:
        a = "可用" if self._deepseek_ready else "不可用"
        b = "可用" if self._qwen_ready else "不可用"
        c = "可用" if self._kimi_ready else "不可用"
        line1 = f"DeepSeek：{a}；千问：{b}；Kimi：{c}"
        line2 = (
            f"资料：{self._cur_source_name}；"
            f"{self._build_question_progress_text()}；"
            f"第{self._cur_round_no}轮"
        )
        self._status_label.setText(line1 + "\n" + line2)

    def _build_question_progress_text(self) -> str:
        total = self._progress_total
        cur_question = str(self._cur_question_no).strip()
        if total > 0:
            if cur_question.isdigit():
                cur = min(max(int(cur_question), 1), total)
            else:
                cur = min(max(self._progress_cur, 0), total)
            return f"选择题 {cur}/{total}题"
        if cur_question.isdigit():
            return f"选择题 {cur_question}/?题"
        return "选择题 统计中"


class _AiImportWorker(QObject):
    progress = pyqtSignal(str)
    question = pyqtSignal(object)
    compare = pyqtSignal(object)
    progress_count = pyqtSignal(int, int)
    done = pyqtSignal(int)
    error = pyqtSignal(str)

    def __init__(self, *, cfg, paths: list[Path], strategy: str, max_question_workers: int) -> None:
        super().__init__()
        self._cfg = cfg
        self._paths = paths
        self._strategy = strategy
        self._max_question_workers = normalize_ai_concurrency(max_question_workers)
        self._stop = False

    def request_stop(self) -> None:
        self._stop = True

    def run(self) -> None:
        try:
            sources: list[tuple[Path, str]] = []
            for p in self._paths:
                if p.suffix.lower() != ".docx":
                    raise RuntimeError(f"当前仅支持 Word 文档导入：{p.name}")
                sources.append((p, read_source_text(p)))

            client = LlmClient(to_llm_config(self._cfg))
            kimi_cfg = load_kimi_config()
            qwen_cfg = load_qwen_config()
            if not kimi_cfg.is_ready() or not qwen_cfg.is_ready():
                raise RuntimeError("请先完成 Kimi 与千问配置并通过可用性测试。")
            kimi_client = LlmClient(to_kimi_llm_config(kimi_cfg))
            qwen_client = LlmClient(to_qwen_llm_config(qwen_cfg))
            result = import_questions_from_sources(
                client=client,
                kimi_client=kimi_client,
                qwen_client=qwen_client,
                client_factory=lambda: LlmClient(to_llm_config(self._cfg)),
                kimi_client_factory=lambda: LlmClient(to_kimi_llm_config(kimi_cfg)),
                qwen_client_factory=lambda: LlmClient(to_qwen_llm_config(qwen_cfg)),
                sources=sources,
                strategy=self._strategy,
                max_question_workers=self._max_question_workers,
                progress_cb=self.progress.emit,
                question_cb=self._emit_question,
                compare_cb=self._emit_compare,
                progress_count_cb=self.progress_count.emit,
                stop_cb=self._should_stop,
            )
            self.done.emit(len(result.questions))
        except Exception as e:
            self.error.emit(str(e))

    def _emit_question(self, q: Question) -> None:
        self.question.emit(q)

    def _emit_compare(self, payload: dict) -> None:
        self.compare.emit(payload)

    def _should_stop(self) -> bool:
        return self._stop
