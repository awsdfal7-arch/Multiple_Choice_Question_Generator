from __future__ import annotations

from pathlib import Path

from docx import Document


def read_source_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".txt":
        return _read_text_with_fallback(path)
    if ext == ".docx":
        doc = Document(str(path))
        parts: list[str] = []
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                parts.append(t)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        t = (p.text or "").strip()
                        if t:
                            parts.append(t)
        return "\n".join(parts) + ("\n" if parts else "")
    return _read_text_with_fallback(path)


def _read_text_with_fallback(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except Exception:
        try:
            return path.read_text(encoding="utf-8", errors="ignore")
        except Exception:
            return path.read_text(encoding="gb18030", errors="ignore")
