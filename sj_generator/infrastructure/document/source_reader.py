from __future__ import annotations

from pathlib import Path

from docx import Document


def read_source_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".docx":
        doc = Document(str(path))
        parts: list[str] = []
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                parts.append(t)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        t = (p.text or "").strip()
                        if t:
                            parts.append(t)
        return "\n".join(parts) + ("\n" if parts else "")
    raise ValueError(f"当前仅支持 Word 文档读取：{path.name}")
